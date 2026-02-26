"""
generate-gapfinder-docx.py

Fills a DOCX template using Phase 1 outputs, then exports a PDF via Microsoft Word (docx2pdf).

Usage (run from project root):
  python scripts/generate-gapfinder-docx.py latexmattress.com.au

Assumes template exists at:
  templates/gapfinder_readiness_template.docx

Inputs:
  data/<domain>/analysis/phase1_inventory.xlsx
  data/<domain>/analysis/unknown_vendors.csv
  data/<domain>/analysis/psi.json

Outputs:
  data/<domain>/report/GapFinder_Readiness_<domain>.docx
  data/<domain>/report/GapFinder_Readiness_<domain>.pdf
"""

import os
import re
import csv
import json
from datetime import datetime
from collections import Counter, defaultdict

from openpyxl import load_workbook
from docx import Document
from docx2pdf import convert


# -----------------------------
# Paths / constants
# -----------------------------
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DATA_DIR = os.path.join(ROOT, "data")
TEMPLATE_PATH = os.path.join(ROOT, "templates", "gapfinder_readiness_template.docx")


# -----------------------------
# Utilities
# -----------------------------
def normalise_domain(inp: str) -> str:
    inp = (inp or "").strip()
    inp = re.sub(r"^https?://", "", inp, flags=re.I)
    inp = inp.split("/")[0]
    return inp.replace("www.", "")

def safe_sheet(wb, names):
    for n in names:
        if n in wb.sheetnames:
            return wb[n]
    return None

def header_map(sheet):
    if not sheet:
        return {}
    m = {}
    for cell in sheet[1]:
        if cell.value is None:
            continue
        m[str(cell.value).strip()] = cell.col_idx
    return m

def cell_str(row, idx):
    if not idx:
        return ""
    v = row[idx - 1].value
    return str(v).strip() if v is not None else ""

def read_unknown_hosts(csv_path, top_n=10):
    if not os.path.exists(csv_path):
        return 0, []
    hosts = []
    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            h = (row.get("Host") or "").strip()
            if h:
                hosts.append(h)
    c = Counter(hosts)
    return len(hosts), [h for h, _ in c.most_common(top_n)]

def normalise_event_name(name: str) -> str:
    n = (name or "").strip()
    if not n:
        return ""
    n_low = n.lower().strip()
    mapping = {
        "pageview": "page_view",
        "page view": "page_view",
        "page_view": "page_view",
        "viewcontent": "view_item",
        "view_content": "view_item",
        "view item": "view_item",
        "view_item": "view_item",
    }
    return mapping.get(n_low, n)

def bullet_lines(items):
    return "\n".join([f"• {i}" for i in items if str(i).strip()])

def yesno(flag: bool) -> str:
    return "Yes" if flag else "Not observed in captured traffic"

def read_json_if_exists(p: str):
    if not os.path.exists(p):
        return None
    try:
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None

def fmt_val(v, suffix=""):
    if v is None or v == "":
        return "N/A"
    return f"{v}{suffix}"


# -----------------------------
# Extraction logic
# -----------------------------
def extract_vendor_presence(tag_sheet):
    if not tag_sheet:
        return {
            "has_gtm": False,
            "has_ga4": False,
            "ad_platforms": [],
            "vendor_count": 0,
            "top_by_category": {}
        }

    hm = header_map(tag_sheet)

    vendors = set()
    categories_seen = []
    ad_platforms = set()
    by_category = defaultdict(list)

    for r in tag_sheet.iter_rows(min_row=2, values_only=False):
        vendor = cell_str(r, hm.get("Vendor"))
        cat = cell_str(r, hm.get("Category"))

        if not vendor:
            continue

        vendors.add(vendor)
        if cat:
            categories_seen.append(cat)

        cat_l = (cat or "").lower()

        if "ads" in cat_l:
            if "Google Ads" in vendor: ad_platforms.add("Google Ads")
            if "Meta" in vendor: ad_platforms.add("Meta")
            if "TikTok" in vendor: ad_platforms.add("TikTok")
            if "Microsoft Ads" in vendor or "UET" in vendor: ad_platforms.add("Microsoft Ads")
            if "Pinterest" in vendor: ad_platforms.add("Pinterest")

        if cat and vendor not in by_category[cat]:
            by_category[cat].append(vendor)

    has_gtm = any("Tag Manager" in (c or "") for c in categories_seen) or any("Google Tag Manager" in v for v in vendors)
    has_ga4 = any(("Google Analytics" in v or "GA4" in v) for v in vendors)

    keep_cats = [
        "Analytics",
        "Ads",
        "Consent/CMP",
        "Session Replay",
        "A/B Testing",
        "Email/SMS",
        "Reviews/UGC",
        "Social Feed",
        "Support/Chat / Lead Capture",
        "Payments",
        "Search / Merchandising",
        "Server-side Tagging / Proxy",
    ]

    top_by_category = {}
    for k in keep_cats:
        vals = by_category.get(k, [])
        if not vals:
            for kk, vv in by_category.items():
                if k.lower() in (kk or "").lower():
                    vals = vv
                    break
        if vals:
            top_by_category[k] = ", ".join(vals[:3])

    return {
        "has_gtm": has_gtm,
        "has_ga4": has_ga4,
        "ad_platforms": sorted(list(ad_platforms)),
        "vendor_count": len(vendors),
        "top_by_category": top_by_category
    }

def extract_event_stats(event_sheet):
    if not event_sheet:
        return {
            "event_count": 0,
            "top_events": [],
            "pct_value": None,
            "pct_currency": None,
            "pct_items": None,
            "pct_txn": None,
            "observed_event_set": set(),
        }

    hm = header_map(event_sheet)

    raw_events = []
    total = 0
    flags = {"HasValue": 0, "HasCurrency": 0, "HasItems": 0, "HasTransactionId": 0}

    for r in event_sheet.iter_rows(min_row=2, values_only=False):
        total += 1
        ev = cell_str(r, hm.get("EventName"))
        if ev:
            raw_events.append(normalise_event_name(ev))

        for k in list(flags.keys()):
            val = cell_str(r, hm.get(k))
            if val.upper() == "Y":
                flags[k] += 1

    c = Counter([e for e in raw_events if e])
    top = [name for name, _ in c.most_common(6)]
    observed_set = set(c.keys())

    def pct(x):
        if total == 0:
            return None
        return round((x / total) * 100)

    return {
        "event_count": total,
        "top_events": top,
        "pct_value": pct(flags["HasValue"]),
        "pct_currency": pct(flags["HasCurrency"]),
        "pct_items": pct(flags["HasItems"]),
        "pct_txn": pct(flags["HasTransactionId"]),
        "observed_event_set": observed_set,
    }


# -----------------------------
# Derived blocks
# -----------------------------
def build_tools_by_function(vendor_info):
    lines = []
    order = [
        "Analytics",
        "Ads",
        "Consent/CMP",
        "Email/SMS",
        "A/B Testing",
        "Session Replay",
        "Reviews/UGC",
        "Support/Chat / Lead Capture",
        "Search / Merchandising",
        "Payments",
        "Server-side Tagging / Proxy",
    ]
    for k in order:
        v = vendor_info["top_by_category"].get(k)
        if v:
            lines.append(f"{k}: {v}")
    return bullet_lines(lines) if lines else "• None observed"

def build_journey_signals(event_info):
    observed = event_info.get("observed_event_set", set())

    def has_any(candidates):
        return any(c in observed for c in candidates)

    browsing = has_any({"page_view", "view_item", "view_item_list"})
    product_interest = has_any({"view_item", "select_item", "ViewContent"})
    cart = has_any({"add_to_cart", "AddToCart"})
    checkout = has_any({"begin_checkout", "InitiateCheckout", "add_shipping_info", "add_payment_info"})
    purchase = has_any({"purchase", "Purchase"})

    lines = [
        f"Browsing signals (e.g. page/product views): {yesno(browsing)}",
        f"Product interest signals (e.g. view_item): {yesno(product_interest)}",
        f"Cart signals (e.g. add_to_cart): {yesno(cart)}",
        f"Checkout signals (e.g. begin_checkout): {yesno(checkout)}",
        f"Purchase signals (e.g. purchase): {yesno(purchase)}",
    ]
    return bullet_lines(lines)

def build_payload_completeness(event_info):
    pv = event_info.get("pct_value")
    pc = event_info.get("pct_currency")
    pi = event_info.get("pct_items")
    pt = event_info.get("pct_txn")

    def fmt(p):
        return "N/A" if p is None else f"{p}%"

    lines = [
        f"Value present (observed): {fmt(pv)}",
        f"Currency present (observed): {fmt(pc)}",
        f"Items present (observed): {fmt(pi)}",
        f"Transaction ID present (observed): {fmt(pt)}",
    ]
    return bullet_lines(lines)

def build_coverage_summary(domain, vendor_info, event_info):
    platforms = ", ".join(vendor_info["ad_platforms"]) if vendor_info["ad_platforms"] else "Not observed in captured traffic"
    lines = [
        f"Domain: {domain}",
        f"Tracking foundation: GTM {yesno(vendor_info['has_gtm'])} • GA4 {yesno(vendor_info['has_ga4'])}",
        f"Paid platforms observed: {platforms}",
        f"Distinct vendors observed: {vendor_info['vendor_count']}",
        f"Total events observed: {event_info.get('event_count', 0)}",
    ]
    return bullet_lines(lines)

def build_attribution_summary():
    lines = [
        "Campaign parameters (UTMs): Not evaluated in this run (probe required)",
        "Platform click IDs (gclid/fbclid/ttclid/wbraid): Not observed in captured traffic",
        "Attribution sensitivity: Reported performance can shift based on configuration choices",
    ]
    return bullet_lines(lines)

def build_unknown_summary(unknown_count, unknown_top):
    if unknown_count == 0:
        return bullet_lines([
            "Unknown vendors observed: 0",
            "Top unknown hosts: None",
        ])
    return bullet_lines([
        f"Unknown vendors observed: {unknown_count}",
        f"Top unknown hosts: {', '.join(unknown_top) if unknown_top else 'None'}",
    ])


# -----------------------------
# DOCX placeholder replacement (run-safe)
# -----------------------------
def _replace_in_paragraph(paragraph, mapping: dict):
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, val in mapping.items():
        token = f"{{{{{key}}}}}"
        if token in new_text:
            new_text = new_text.replace(token, str(val))

    if new_text == full_text:
        return

    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""

def _replace_in_cell(cell, mapping: dict):
    for p in cell.paragraphs:
        _replace_in_paragraph(p, mapping)

def replace_placeholders_in_doc(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _replace_in_cell(cell, mapping)

    for section in doc.sections:
        header = section.header
        footer = section.footer

        for p in header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in header.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)

        for p in footer.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)


# -----------------------------
# PSI → placeholder mapping (home target)
# -----------------------------
def apply_psi(mapping: dict, psi: dict):
    """
    Populates v1-style CWV placeholders for the DOCX table.
    Uses the "home" target by default.
    """
    home = (psi or {}).get("targets", {}).get("home", {}) if psi else {}
    m = (home or {}).get("mobile", {})
    d = (home or {}).get("desktop", {})

    # Match your DOCX placeholders exactly (no renaming)
    mapping.update({
        "MobileCLS": fmt_val(m.get("cls")),
        "MobileLCP": fmt_val(m.get("lcp_s"), "s"),
        "MobileFCP": fmt_val(m.get("fcp_s"), "s"),
        "MobileTTI": fmt_val(m.get("tti_s"), "s"),
        "MobileTBT": fmt_val(m.get("tbt_ms"), "ms"),
        "MobileOverall": fmt_val(m.get("performance")),

        "DesktopCLS": fmt_val(d.get("cls")),
        "DesktopLCP": fmt_val(d.get("lcp_s"), "s"),
        "DesktopFCP": fmt_val(d.get("fcp_s"), "s"),
        "DesktopTTI": fmt_val(d.get("tti_s"), "s"),
        "DesktopTBT": fmt_val(d.get("tbt_ms"), "ms"),
        "DesktopOverall": fmt_val(d.get("performance")),
    })


# -----------------------------
# Main
# -----------------------------
def main(domain_input: str):
    domain = normalise_domain(domain_input)

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    analysis_dir = os.path.join(DATA_DIR, domain, "analysis")
    xlsx_path = os.path.join(analysis_dir, "phase1_inventory.xlsx")
    unknown_path = os.path.join(analysis_dir, "unknown_vendors.csv")
    psi_path = os.path.join(analysis_dir, "psi.json")

    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"Missing workbook: {xlsx_path}")

    wb = load_workbook(xlsx_path, data_only=True)

    tag_sheet = safe_sheet(wb, ["baseline_tag_inventory", "tag_inventory", "Tag Inventory"])
    event_sheet = safe_sheet(wb, ["baseline_event_inventory", "event_inventory", "Event Inventory"])

    vendor_info = extract_vendor_presence(tag_sheet)
    event_info = extract_event_stats(event_sheet)
    unknown_count, unknown_top = read_unknown_hosts(unknown_path, top_n=10)

    mapping = {
        "website": domain,
        "GeneratedAt": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "HasGTM": "Yes" if vendor_info["has_gtm"] else "Not observed in captured traffic",
        "HasGA4": "Yes" if vendor_info["has_ga4"] else "Not observed in captured traffic",
        "AdPlatforms": ", ".join(vendor_info["ad_platforms"]) if vendor_info["ad_platforms"] else "Not observed in captured traffic",
        "VendorCount": str(vendor_info["vendor_count"]),
        "TopEvents": ", ".join(event_info["top_events"]) if event_info["top_events"] else "None observed",
        "EventCount": str(event_info["event_count"]),
        "PctHasValue": "N/A" if event_info["pct_value"] is None else f"{event_info['pct_value']}%",
        "PctHasCurrency": "N/A" if event_info["pct_currency"] is None else f"{event_info['pct_currency']}%",
        "PctHasItems": "N/A" if event_info["pct_items"] is None else f"{event_info['pct_items']}%",
        "PctHasTransactionId": "N/A" if event_info["pct_txn"] is None else f"{event_info['pct_txn']}%",
        "UTMsObserved": "Inconclusive from website signals alone",
        "UTMsPersist": "Inconclusive from website signals alone",
        "ClickIdsObserved": "Not observed in captured traffic",
        "UnknownCount": str(unknown_count),
        "TopUnknownHosts": ", ".join(unknown_top) if unknown_top else "None",

        # PSI placeholders default (if psi.json missing)
        "MobileCLS": "N/A",
        "MobileLCP": "N/A",
        "MobileFCP": "N/A",
        "MobileTTI": "N/A",
        "MobileTBT": "N/A",
        "MobileOverall": "N/A",
        "DesktopCLS": "N/A",
        "DesktopLCP": "N/A",
        "DesktopFCP": "N/A",
        "DesktopTTI": "N/A",
        "DesktopTBT": "N/A",
        "DesktopOverall": "N/A",
    }

    mapping["CoverageSummary"] = build_coverage_summary(domain, vendor_info, event_info)
    mapping["ToolsByFunction"] = build_tools_by_function(vendor_info)
    mapping["JourneySignals"] = build_journey_signals(event_info)
    mapping["PayloadCompleteness"] = build_payload_completeness(event_info)
    mapping["AttributionSummary"] = build_attribution_summary()
    mapping["UnknownSummary"] = build_unknown_summary(unknown_count, unknown_top)

    parts = []
    for cat, vendors in vendor_info["top_by_category"].items():
        parts.append(f"{cat}: {vendors}")
    mapping["TopVendorsByCategory"] = " | ".join(parts) if parts else "None observed"

    # Apply PSI if available
    psi = read_json_if_exists(psi_path)
    if psi:
        apply_psi(mapping, psi)

    doc = Document(TEMPLATE_PATH)
    replace_placeholders_in_doc(doc, mapping)

    report_dir = os.path.join(DATA_DIR, domain, "report")
    os.makedirs(report_dir, exist_ok=True)

    out_docx = os.path.join(report_dir, f"GapFinder_Readiness_{domain}.docx")
    doc.save(out_docx)

    out_pdf = os.path.splitext(out_docx)[0] + ".pdf"
    try:
        convert(out_docx, out_pdf)
    except Exception as e:
        raise RuntimeError(
            "DOCX saved, but PDF conversion failed. "
            "Make sure Microsoft Word is installed and the DOCX is not open.\n"
            f"Details: {e}"
        )

    print(f"[OK] Wrote DOCX: {out_docx}")
    print(f"[OK] Wrote PDF:  {out_pdf}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python scripts/generate-gapfinder-docx.py <domain>")
        raise SystemExit(1)
    main(sys.argv[1])